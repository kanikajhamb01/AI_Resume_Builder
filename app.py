import streamlit as st
import google.generativeai as genai
from docx import Document
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
import os
import uuid

# --- Sidebar: Gemini API key ---
st.sidebar.title("🔐 Gemini API Settings")
gemini_api_key = st.sidebar.text_input("Enter your Gemini API Key", type="password")
if not gemini_api_key:
    st.sidebar.warning("Please enter your Gemini API key.")
    st.stop()

# Configure Gemini
genai.configure(api_key=gemini_api_key)

# --- App title ---
st.title("🤖 AI Resume Builder (Powered by Gemini)")

# --- Form UI ---
with st.form("resume_form"):
    name = st.text_input("Full Name")
    email = st.text_input("Email")
    education = st.text_input("Education")
    skills = st.text_input("Skills (comma separated)")

    st.markdown("### Work Experience")
    role = st.text_input("Job Role")
    company = st.text_input("Company Name")
    description = st.text_area("Job Description")

    st.markdown("### Project")
    project_title = st.text_input("Project Title")
    project_desc = st.text_area("Project Description")

    export_format = st.selectbox("Choose Resume Format", ["DOCX", "PDF"])
    template_choice = st.selectbox("Choose Resume Template (for PDF only)", ["Modern", "Classic", "Minimal"])

    submit = st.form_submit_button("Generate Resume")

# --- Gemini bullet generators ---
def generate_experience_bullets(role, company, description):
    prompt = f"""
You are a professional resume writer.

Write 2 quantified, professional, ATS-optimized bullet points for the following experience:

Role: {role}
Company: {company}
Description: {description}

Respond with only the bullet points. Each bullet should start with •
"""
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"⚠️ Error generating experience bullets: {e}"

def generate_project_bullet(title, description):
    prompt = f"""
You are a professional resume writer.

Write 1 strong, quantified, and concise bullet point for the following project:

Title: {title}
Description: {description}

Respond with only the bullet point. Start it with •
"""
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"⚠️ Error generating project bullet: {e}"

# --- Resume creation ---
def create_resume(name, email, education, skills, job, company, job_desc, proj_title, proj_desc, export_format, template_choice):
    experience_bullets = generate_experience_bullets(job, company, job_desc)
    exp_bullets = [b.strip("• ").strip() for b in experience_bullets.strip().split("•") if b.strip()]
    project_bullet = generate_project_bullet(proj_title, proj_desc).strip("• ").strip()
    skills_list = [s.strip() for s in skills.split(",")]

    if export_format == "DOCX":
        doc = Document()
        doc.add_heading(name, 0)
        doc.add_paragraph(email)

        doc.add_heading("Education", level=1)
        doc.add_paragraph(education)

        doc.add_heading("Skills", level=1)
        for skill in skills_list:
            doc.add_paragraph(skill, style='List Bullet')

        doc.add_heading("Experience", level=1)
        doc.add_paragraph(f"{job} at {company}", style='List Bullet')
        for bullet in exp_bullets:
            doc.add_paragraph(bullet, style='List Bullet')

        doc.add_heading("Projects", level=1)
        doc.add_paragraph(proj_title, style='List Bullet')
        doc.add_paragraph(project_bullet, style='List Bullet')

        filename = f"AI_Resume_{uuid.uuid4().hex[:6]}.docx"
        doc.save(filename)
        return filename, "docx"

    else:
        env = Environment(loader=FileSystemLoader("templates"))
        template_file = {
            "Modern": "modern.html",
            "Classic": "classic.html",
            "Minimal": "minimal.html"
        }.get(template_choice, "modern.html")

        try:
            template = env.get_template(template_file)
            html_out = template.render(
                name=name,
                email=email,
                education=education,
                skills=skills_list,
                role=job,
                company=company,
                exp_bullets=exp_bullets,
                project_title=proj_title,
                project_bullet=project_bullet
            )
            filename = f"AI_Resume_{uuid.uuid4().hex[:6]}.pdf"
            HTML(string=html_out).write_pdf(filename)
            return filename, "pdf"
        except Exception as e:
            st.error(f"❌ PDF Generation Error: {e}")
            st.stop()

# --- Generate and Download Resume ---
if submit:
    if not all([name, email, education, skills, role, company, description, project_title, project_desc]):
        st.error("❗ Please fill in all fields.")
    else:
        with st.spinner("Generating your resume with Gemini..."):
            resume_file, file_type = create_resume(
                name, email, education, skills,
                role, company, description,
                project_title, project_desc,
                export_format, template_choice
            )
        with open(resume_file, "rb") as file:
            st.success("✅ Resume Generated Successfully!")
            st.download_button(
                label=f"📥 Download Resume ({file_type.upper()})",
                data=file,
                file_name=resume_file,
                mime="application/pdf" if file_type == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        os.remove(resume_file)







